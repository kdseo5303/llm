import pandas as pd
import io
import os
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import openpyxl
from docx import Document
import PyPDF2
import aiofiles
from ..models.chat import KnowledgeDocument

class FileProcessor:
    """Service for processing various file formats and extracting text content."""
    
    def __init__(self):
        self.supported_extensions = {
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv,
            '.txt': self._process_text,
            '.docx': self._process_docx,
            '.pdf': self._process_pdf
        }
    
    async def process_file(self, file_content: bytes, filename: str, category: str) -> List[KnowledgeDocument]:
        """
        Process uploaded file and extract knowledge documents.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            category: Document category (pre-production, production, post-production)
            
        Returns:
            List of KnowledgeDocument objects
        """
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Get the appropriate processor
        processor = self.supported_extensions[file_extension]
        
        try:
            # Process the file
            documents = await processor(file_content, filename, category)
            return documents
        except Exception as e:
            raise Exception(f"Error processing file {filename}: {str(e)}")
    
    async def _process_excel(self, file_content: bytes, filename: str, category: str) -> List[KnowledgeDocument]:
        """Process Excel files (.xlsx, .xls)."""
        documents = []
        
        try:
            # Read Excel file
            excel_file = io.BytesIO(file_content)
            
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            
            for sheet_name, df in excel_data.items():
                if df.empty:
                    continue
                
                # Convert DataFrame to text
                text_content = self._dataframe_to_text(df, sheet_name)
                
                # Create document for each sheet
                doc = KnowledgeDocument(
                    id=f"{filename}_{sheet_name}",
                    title=f"{Path(filename).stem} - {sheet_name}",
                    content=text_content,
                    category=category,
                    tags=[category, "excel", sheet_name.lower()]
                )
                
                documents.append(doc)
            
            # If no sheets or all empty, create one document from the file
            if not documents:
                # Try to read as single sheet
                excel_file.seek(0)
                df = pd.read_excel(excel_file, engine='openpyxl')
                if not df.empty:
                    text_content = self._dataframe_to_text(df, "Sheet1")
                    doc = KnowledgeDocument(
                        id=filename,
                        title=Path(filename).stem,
                        content=text_content,
                        category=category,
                        tags=[category, "excel"]
                    )
                    documents.append(doc)
            
        except Exception as e:
            raise Exception(f"Error processing Excel file: {str(e)}")
        
        return documents
    
    async def _process_csv(self, file_content: bytes, filename: str, category: str) -> List[KnowledgeDocument]:
        """Process CSV files."""
        try:
            # Read CSV content
            csv_content = file_content.decode('utf-8')
            df = pd.read_csv(io.StringIO(csv_content))
            
            # Convert to text
            text_content = self._dataframe_to_text(df, "CSV Data")
            
            doc = KnowledgeDocument(
                id=filename,
                title=Path(filename).stem,
                content=text_content,
                category=category,
                tags=[category, "csv"]
            )
            
            return [doc]
            
        except Exception as e:
            raise Exception(f"Error processing CSV file: {str(e)}")
    
    async def _process_text(self, file_content: bytes, filename: str, category: str) -> List[KnowledgeDocument]:
        """Process text files."""
        try:
            # Decode text content
            text_content = file_content.decode('utf-8')
            
            doc = KnowledgeDocument(
                id=filename,
                title=Path(filename).stem,
                content=text_content,
                category=category,
                tags=[category, "text"]
            )
            
            return [doc]
            
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    async def _process_docx(self, file_content: bytes, filename: str, category: str) -> List[KnowledgeDocument]:
        """Process Word documents (.docx)."""
        try:
            # Read Word document
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text.strip() + "\n\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_content += row_text + "\n"
                text_content += "\n"
            
            doc_obj = KnowledgeDocument(
                id=filename,
                title=Path(filename).stem,
                content=text_content.strip(),
                category=category,
                tags=[category, "word"]
            )
            
            return [doc_obj]
            
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    async def _process_pdf(self, file_content: bytes, filename: str, category: str) -> List[KnowledgeDocument]:
        """Process PDF files."""
        try:
            # Read PDF content
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content += f"--- Page {page_num + 1} ---\n{page_text.strip()}\n\n"
            
            doc = KnowledgeDocument(
                id=filename,
                title=Path(filename).stem,
                content=text_content.strip(),
                category=category,
                tags=[category, "pdf"]
            )
            
            return [doc]
            
        except Exception as e:
            raise Exception(f"Error processing PDF file: {str(e)}")
    
    def _dataframe_to_text(self, df: pd.DataFrame, sheet_name: str) -> str:
        """Convert pandas DataFrame to readable text format."""
        if df.empty:
            return f"Sheet: {sheet_name}\nNo data found."
        
        # Get column names
        columns = df.columns.tolist()
        
        # Start building text content
        text_content = f"Sheet: {sheet_name}\n"
        text_content += f"Columns: {', '.join(columns)}\n"
        text_content += f"Total Rows: {len(df)}\n\n"
        
        # Add data rows
        for index, row in df.iterrows():
            row_text = f"Row {index + 1}: "
            row_data = []
            
            for col in columns:
                value = row[col]
                if pd.notna(value):  # Check for NaN values
                    row_data.append(f"{col}: {value}")
            
            if row_data:
                text_content += row_text + " | ".join(row_data) + "\n"
        
        return text_content
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_extensions.keys())
    
    def get_file_size_limit(self) -> int:
        """Get maximum file size limit in bytes (50MB)."""
        return 50 * 1024 * 1024  # 50MB
    
    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """
        Validate uploaded file.
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        if file_size > self.get_file_size_limit():
            return False, f"File size {file_size / (1024*1024):.1f}MB exceeds limit of 50MB"
        
        # Check file extension
        file_extension = Path(filename).suffix.lower()
        if file_extension not in self.supported_extensions:
            return False, f"Unsupported file type: {file_extension}. Supported: {', '.join(self.get_supported_formats())}"
        
        return True, "" 